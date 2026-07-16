import io

import docx as docx_lib
import pytest

from app.converters import ConversionError
from app.converters.docx import convert_docx


def _docx_with_heading() -> bytes:
    d = docx_lib.Document()
    d.add_heading("사업 개요", level=1)
    d.add_paragraph("본문 문단입니다.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "항목"
    t.rows[0].cells[1].text = "금액 1억원"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_headings_become_markdown():
    r = convert_docx(_docx_with_heading())
    assert "# 사업 개요" in r.text
    assert "본문 문단입니다." in r.text


def test_table_flattened_with_warning():
    r = convert_docx(_docx_with_heading())
    assert "금액 1억원" in r.text
    assert any("표" in w for w in r.warnings)


def test_empty_docx_rejected():
    d = docx_lib.Document()
    buf = io.BytesIO()
    d.save(buf)
    with pytest.raises(ConversionError):
        convert_docx(buf.getvalue())
